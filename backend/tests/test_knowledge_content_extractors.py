from pathlib import Path

import pytest

from app.services.knowledge_content import (
    ExtractedBlock,
    KnowledgeContentError,
    MAX_CHUNK_CHARACTERS,
    build_document_chunks,
    extract_document_blocks,
)
from app.services import knowledge_content


def test_pdf_docx_and_xlsx_extractors_preserve_source_locations(tmp_path):
    from docx import Document
    from openpyxl import Workbook
    from reportlab.pdfgen.canvas import Canvas

    pdf_path = tmp_path / "source.pdf"
    canvas = Canvas(str(pdf_path))
    canvas.drawString(72, 720, "V10 ambient light brightness is standard")
    canvas.showPage()
    canvas.drawString(72, 720, "Needle gain is tender support")
    canvas.save()

    docx_path = tmp_path / "release-note.docx"
    document = Document()
    document.add_heading("Release Note", level=1)
    document.add_paragraph("Needle gain adjustment is available.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Feature"
    table.cell(0, 1).text = "Status"
    document.save(docx_path)

    xlsx_path = tmp_path / "registration.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "CN"
    worksheet.append(["Model", "Probe", "Status"])
    worksheet.append(["VINNO 10", "X4-9E", "Registered"])
    workbook.save(xlsx_path)

    pdf_blocks = extract_document_blocks(pdf_path, "application/pdf")
    docx_blocks = extract_document_blocks(
        docx_path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    xlsx_blocks = extract_document_blocks(
        xlsx_path,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    assert [block.source_ref for block in pdf_blocks] == ["第1页", "第2页"]
    assert pdf_blocks[0].page_number == 1
    assert any(block.section_name == "Release Note" for block in docx_blocks)
    assert any(block.source_ref == "表格1 第1行" for block in docx_blocks)
    assert [block.source_ref for block in xlsx_blocks] == ["CN!第1行", "CN!第2行"]


def test_text_extraction_long_chunking_and_unsupported_type(tmp_path):
    text_path = tmp_path / "notes.txt"
    text_path.write_text(" first line \n\n second   line ", encoding="utf-8")

    blocks = extract_document_blocks(text_path, "text/plain")
    chunks = build_document_chunks(blocks)
    long_chunks = build_document_chunks(
        [ExtractedBlock(text="A" * (MAX_CHUNK_CHARACTERS * 2 + 20), source_ref="全文")]
    )

    assert chunks[0]["content"] == "first line\nsecond line"
    assert chunks[0]["source_ref"] == "全文"
    assert len(long_chunks) == 3
    assert all(len(chunk["content"]) <= MAX_CHUNK_CHARACTERS for chunk in long_chunks)

    unsupported = tmp_path / "source.bin"
    unsupported.write_bytes(b"binary")
    with pytest.raises(KnowledgeContentError, match="暂不支持"):
        extract_document_blocks(unsupported, "application/octet-stream")


def test_empty_text_produces_no_blocks_or_chunks(tmp_path):
    empty_path = tmp_path / "empty.txt"
    empty_path.write_text(" \n\t ", encoding="utf-8")

    assert extract_document_blocks(empty_path, "text/plain") == []
    assert build_document_chunks([ExtractedBlock(text="", source_ref="全文")]) == []


def test_image_only_pdf_falls_back_to_ocr(tmp_path, monkeypatch):
    from pypdf import PdfWriter

    pdf_path = tmp_path / "scanned-registration.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    with pdf_path.open("wb") as stream:
        writer.write(stream)

    expected = [
        ExtractedBlock(
            text="注册证编号：苏械注准20232061322",
            source_ref="第1页（OCR）",
            page_number=1,
        )
    ]
    monkeypatch.setattr(
        knowledge_content,
        "_ocr_pdf_blocks",
        lambda path: expected,
        raising=False,
    )

    assert extract_document_blocks(pdf_path, "application/pdf") == expected
