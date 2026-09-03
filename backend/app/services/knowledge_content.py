"""Extract controlled source text and retrieve unconfirmed candidate evidence."""

from __future__ import annotations

import asyncio
from dataclasses import dataclass
import hashlib
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
import unicodedata

from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from app.services.knowledge_qa import normalize_question


EXTRACTOR_VERSION = "2"
MAX_CHUNK_CHARACTERS = 1200
CHUNK_OVERLAP_CHARACTERS = 120
MAX_CANDIDATE_EXCERPT = 800
_QUESTION_FILLERS = (
    "是不是",
    "是否",
    "这个",
    "该功能",
    "请问",
    "的吗",
    "吗",
    "呢",
    "有无",
    "有没有",
)
_CONFIRMED_QUERY_EQUIVALENCES = (
    (
        "环境光自动调节亮度",
        "环境光线自动调节亮度",
        "实时显示屏亮度调节",
        "显示屏亮度调节",
    ),
)


class KnowledgeContentError(ValueError):
    """Raised when a controlled document cannot be extracted."""


class KnowledgeDocumentMissingError(KnowledgeContentError):
    """Raised when the registered original file is unavailable."""


@dataclass(frozen=True)
class ExtractedBlock:
    text: str
    source_ref: str
    page_number: int | None = None
    section_name: str | None = None


def _clean_text(value: str) -> str:
    normalized = unicodedata.normalize("NFKC", str(value or ""))
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in normalized.splitlines()]
    return "\n".join(line for line in lines if line)


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _pdf_blocks(path: Path) -> list[ExtractedBlock]:
    from pypdf import PdfReader

    reader = PdfReader(path)
    blocks: list[ExtractedBlock] = []
    for page_index, page in enumerate(reader.pages, start=1):
        page_text = _clean_text(page.extract_text() or "")
        if page_text:
            blocks.append(
                ExtractedBlock(
                    text=page_text,
                    source_ref=f"第{page_index}页",
                    page_number=page_index,
                )
            )
    return blocks or _ocr_pdf_blocks(path)


def _ocr_pdf_blocks(path: Path) -> list[ExtractedBlock]:
    pdftoppm = shutil.which("pdftoppm")
    tesseract = shutil.which("tesseract")
    if not pdftoppm or not tesseract:
        return []

    with tempfile.TemporaryDirectory(prefix="vinno-pdf-ocr-") as directory:
        page_prefix = Path(directory) / "page"
        try:
            subprocess.run(
                [pdftoppm, "-png", "-r", "250", str(path), str(page_prefix)],
                check=True,
                capture_output=True,
                timeout=120,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            raise KnowledgeContentError("图片型 PDF 页面渲染失败") from exc

        page_paths = sorted(
            Path(directory).glob("page-*.png"),
            key=lambda item: int(item.stem.rsplit("-", 1)[-1]),
        )
        blocks: list[ExtractedBlock] = []
        for page_number, page_path in enumerate(page_paths, start=1):
            try:
                result = subprocess.run(
                    [
                        tesseract,
                        str(page_path),
                        "stdout",
                        "-l",
                        "chi_sim+eng",
                        "--psm",
                        "3",
                    ],
                    check=True,
                    capture_output=True,
                    text=True,
                    timeout=120,
                )
            except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
                raise KnowledgeContentError(
                    f"图片型 PDF 第{page_number}页 OCR 失败"
                ) from exc
            page_text = _clean_text(result.stdout)
            if page_text:
                blocks.append(
                    ExtractedBlock(
                        text=page_text,
                        source_ref=f"第{page_number}页（OCR）",
                        page_number=page_number,
                    )
                )
        return blocks


def _docx_blocks(path: Path) -> list[ExtractedBlock]:
    from docx import Document

    document = Document(path)
    blocks: list[ExtractedBlock] = []
    section_name: str | None = None
    paragraph_number = 0
    for paragraph in document.paragraphs:
        content = _clean_text(paragraph.text)
        if not content:
            continue
        paragraph_number += 1
        style_name = str(paragraph.style.name or "") if paragraph.style else ""
        if style_name.casefold().startswith("heading") or style_name.startswith("标题"):
            section_name = content
        blocks.append(
            ExtractedBlock(
                text=content,
                source_ref=section_name or f"段落{paragraph_number}",
                section_name=section_name,
            )
        )
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [_clean_text(cell.text) for cell in row.cells]
            content = " | ".join(value for value in values if value)
            if content:
                blocks.append(
                    ExtractedBlock(
                        text=content,
                        source_ref=f"表格{table_index} 第{row_index}行",
                        section_name=f"表格{table_index}",
                    )
                )
    return blocks


def _xlsx_blocks(path: Path) -> list[ExtractedBlock]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    blocks: list[ExtractedBlock] = []
    try:
        for worksheet in workbook.worksheets:
            for row_index, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
                values = [_clean_text(str(value)) for value in row if value is not None]
                content = " | ".join(value for value in values if value)
                if content:
                    blocks.append(
                        ExtractedBlock(
                            text=content,
                            source_ref=f"{worksheet.title}!第{row_index}行",
                            section_name=worksheet.title,
                        )
                    )
    finally:
        workbook.close()
    return blocks


def extract_document_blocks(path: Path, mime_type: str | None) -> list[ExtractedBlock]:
    suffix = path.suffix.casefold()
    mime = str(mime_type or "").casefold()
    if mime == "application/pdf" or suffix == ".pdf":
        return _pdf_blocks(path)
    if "wordprocessingml.document" in mime or suffix == ".docx":
        return _docx_blocks(path)
    if "spreadsheetml.sheet" in mime or suffix == ".xlsx":
        return _xlsx_blocks(path)
    if mime.startswith("text/") or suffix in {".txt", ".md", ".csv"}:
        content = _clean_text(path.read_text(encoding="utf-8", errors="replace"))
        return [ExtractedBlock(text=content, source_ref="全文")] if content else []
    raise KnowledgeContentError(f"暂不支持提取此文件类型：{mime_type or suffix or '未知'}")


def _split_long_text(content: str) -> list[str]:
    if len(content) <= MAX_CHUNK_CHARACTERS:
        return [content]
    paragraphs = [part.strip() for part in content.split("\n") if part.strip()]
    if not paragraphs:
        paragraphs = [content]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(paragraph) > MAX_CHUNK_CHARACTERS:
            if current:
                chunks.append(current)
                current = ""
            start = 0
            while start < len(paragraph):
                end = min(len(paragraph), start + MAX_CHUNK_CHARACTERS)
                chunks.append(paragraph[start:end])
                if end == len(paragraph):
                    break
                start = max(start + 1, end - CHUNK_OVERLAP_CHARACTERS)
            continue
        combined = f"{current}\n{paragraph}".strip() if current else paragraph
        if len(combined) <= MAX_CHUNK_CHARACTERS:
            current = combined
        else:
            chunks.append(current)
            current = paragraph
    if current:
        chunks.append(current)
    return chunks


def build_document_chunks(blocks: list[ExtractedBlock]) -> list[dict]:
    chunks: list[dict] = []
    for block in blocks:
        cleaned = _clean_text(block.text)
        if not cleaned:
            continue
        for content in _split_long_text(cleaned):
            chunks.append(
                {
                    "chunk_index": len(chunks),
                    "page_number": block.page_number,
                    "section_name": block.section_name,
                    "source_ref": block.source_ref,
                    "content": content,
                    "normalized_content": normalize_question(content),
                    "content_hash": hashlib.sha256(content.encode("utf-8")).hexdigest(),
                }
            )
    return chunks


async def extract_registered_document(
    session: AsyncSession,
    document_id: int,
    *,
    force: bool = False,
) -> dict | None:
    document_result = await session.execute(
        text(
            """
            SELECT id, file_path, mime_type
            FROM knowledge_documents
            WHERE id = :document_id AND source_status = 'active'
            """
        ),
        {"document_id": document_id},
    )
    document = document_result.one_or_none()
    if document is None:
        return None
    path = Path(document.file_path)
    if not path.is_absolute() or not path.is_file():
        raise KnowledgeDocumentMissingError("原文件不存在或尚未同步")

    source_sha256 = await asyncio.to_thread(_file_sha256, path)
    current_result = await session.execute(
        text(
            """
            SELECT extractor_version, source_sha256, status, chunk_count, extracted_at
            FROM knowledge_document_extractions
            WHERE document_id = :document_id
            """
        ),
        {"document_id": document_id},
    )
    current = current_result.one_or_none()
    if (
        not force
        and current is not None
        and current.status == "completed"
        and current.extractor_version == EXTRACTOR_VERSION
        and current.source_sha256 == source_sha256
    ):
        return {
            "document_id": document_id,
            "status": current.status,
            "chunk_count": int(current.chunk_count),
            "extracted_at": current.extracted_at,
            "extractor_version": current.extractor_version,
            "reused": True,
        }

    blocks = await asyncio.to_thread(extract_document_blocks, path, document.mime_type)
    chunks = build_document_chunks(blocks)
    if not chunks:
        raise KnowledgeContentError("未从原文件中提取到可检索正文")

    await session.execute(
        text("DELETE FROM knowledge_document_chunks WHERE document_id = :document_id"),
        {"document_id": document_id},
    )
    for chunk in chunks:
        await session.execute(
            text(
                """
                INSERT INTO knowledge_document_chunks (
                    document_id, chunk_index, page_number, section_name,
                    source_ref, content, normalized_content, content_hash
                ) VALUES (
                    :document_id, :chunk_index, :page_number, :section_name,
                    :source_ref, :content, :normalized_content, :content_hash
                )
                """
            ),
            {"document_id": document_id, **chunk},
        )
    await session.execute(
        text(
            """
            INSERT INTO knowledge_document_extractions (
                document_id, extractor_version, source_sha256, status,
                chunk_count, error_message, extracted_at, updated_at
            ) VALUES (
                :document_id, :extractor_version, :source_sha256, 'completed',
                :chunk_count, NULL, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP
            )
            ON CONFLICT(document_id) DO UPDATE SET
                extractor_version = excluded.extractor_version,
                source_sha256 = excluded.source_sha256,
                status = 'completed',
                chunk_count = excluded.chunk_count,
                error_message = NULL,
                extracted_at = CURRENT_TIMESTAMP,
                updated_at = CURRENT_TIMESTAMP
            """
        ),
        {
            "document_id": document_id,
            "extractor_version": EXTRACTOR_VERSION,
            "source_sha256": source_sha256,
            "chunk_count": len(chunks),
        },
    )
    await session.commit()
    result = await session.execute(
        text(
            """
            SELECT status, chunk_count, extracted_at, extractor_version
            FROM knowledge_document_extractions WHERE document_id = :document_id
            """
        ),
        {"document_id": document_id},
    )
    extraction = result.one()
    return {
        "document_id": document_id,
        "status": extraction.status,
        "chunk_count": int(extraction.chunk_count),
        "extracted_at": extraction.extracted_at,
        "extractor_version": extraction.extractor_version,
        "reused": False,
    }


def _search_core(value: str) -> str:
    core = normalize_question(value)
    for filler in _QUESTION_FILLERS:
        core = core.replace(normalize_question(filler), "")
    return core


def _bigrams(value: str) -> set[str]:
    return {value[index : index + 2] for index in range(max(0, len(value) - 1))}


def _query_forms(question: str) -> list[str]:
    primary = _search_core(question)
    forms = [primary]
    normalized_question = normalize_question(question)
    for group in _CONFIRMED_QUERY_EQUIVALENCES:
        normalized_group = [normalize_question(item) for item in group]
        if any(item in normalized_question for item in normalized_group):
            forms.extend(item for item in normalized_group if item not in forms)
    return forms


def _candidate_prefilter_pairs(question: str) -> tuple[str, ...]:
    pairs = {
        pair
        for form in _query_forms(question)
        for pair in _bigrams(form)
    }
    non_numeric_pairs = {pair for pair in pairs if not pair.isdigit()}
    return tuple(sorted(non_numeric_pairs or pairs))


def _canonical_identifiers(value: str) -> set[str]:
    normalized = normalize_question(value)
    identifiers = set(re.findall(r"(?:vinno|v)\d+[a-z0-9]*|\d{6,}", normalized))
    return {re.sub(r"^vinno", "v", identifier) for identifier in identifiers}


def _pair_coverage(query_core: str, content_core: str) -> float:
    query_pairs = _bigrams(query_core)
    if not query_pairs:
        return 1.0 if query_core and query_core in content_core else 0.0
    if query_core in content_core:
        return 1.0
    return sum(pair in content_core for pair in query_pairs) / len(query_pairs)


def _candidate_score(question: str, content: str, document_context: str) -> float:
    query_forms = _query_forms(question)
    content_core = normalize_question(content)
    context_core = normalize_question(document_context)
    if not query_forms[0] or not content_core:
        return 0.0
    coverage = max(_pair_coverage(form, content_core) for form in query_forms)
    identifiers = _canonical_identifiers(question)
    combined_context = re.sub(r"vinno(?=\d)", "v", f"{content_core}{context_core}")
    if identifiers and not all(identifier in combined_context for identifier in identifiers):
        coverage *= 0.35
    intent_terms = {
        term for term in ("标配", "选配", "招标", "未注册", "注册", "支持", "参数")
        if term in question
    }
    if intent_terms:
        matched_intents = sum(term in content for term in intent_terms)
        coverage = 0.85 * coverage + 0.15 * (matched_intents / len(intent_terms))
    registration_intent = any(term in question for term in ("注册", "湘证", "苏证"))
    registration_source = any(
        marker in context_core
        for marker in ("registrationcertificate", "registrationdifference")
    )
    if registration_intent:
        coverage = (
            min(1.0, coverage + 0.25)
            if registration_source
            else coverage * 0.65
        )
    return round(min(1.0, coverage), 4)


def _candidate_excerpt(content: str, question: str) -> str:
    if len(content) <= MAX_CANDIDATE_EXCERPT:
        return content
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    if not lines:
        return content[:MAX_CANDIDATE_EXCERPT]
    forms = _query_forms(question)
    best_index = max(
        range(len(lines)),
        key=lambda index: max(
            _pair_coverage(form, normalize_question(lines[index])) for form in forms
        ),
    )
    selected = lines[best_index]
    left = best_index - 1
    right = best_index + 1
    while len(selected) < MAX_CANDIDATE_EXCERPT and (left >= 0 or right < len(lines)):
        if left >= 0:
            candidate = f"{lines[left]}\n{selected}"
            if len(candidate) <= MAX_CANDIDATE_EXCERPT:
                selected = candidate
            left -= 1
        if right < len(lines):
            candidate = f"{selected}\n{lines[right]}"
            if len(candidate) <= MAX_CANDIDATE_EXCERPT:
                selected = candidate
            right += 1
    return selected[-MAX_CANDIDATE_EXCERPT:]


async def find_candidate_evidence(
    session: AsyncSession,
    question: str,
    *,
    limit: int = 5,
) -> list[dict]:
    prefilter_pairs = _candidate_prefilter_pairs(question)
    params = {
        f"prefilter_{index}": f"%{pair}%"
        for index, pair in enumerate(prefilter_pairs)
    }
    prefilter_sql = ""
    if params:
        prefilter_sql = "AND (" + " OR ".join(
            f"chunk.normalized_content LIKE :{name}" for name in params
        ) + ")"
    result = await session.execute(
        text(
            f"""
            SELECT chunk.id AS chunk_id, chunk.document_id, document.title,
                   document.document_type, document.product_series,
                   chunk.source_ref, chunk.page_number, chunk.content
            FROM knowledge_document_chunks chunk
            JOIN knowledge_documents document ON document.id = chunk.document_id
            JOIN knowledge_document_extractions extraction
              ON extraction.document_id = document.id
             AND extraction.status = 'completed'
            WHERE document.source_status = 'active'
              {prefilter_sql}
            ORDER BY chunk.document_id, chunk.chunk_index
            """
        ),
        params,
    )
    candidates: list[dict] = []
    for row in result:
        score = _candidate_score(
            question,
            row.content,
            f"{row.document_type} {row.title} {row.product_series or ''}",
        )
        if score < 0.3:
            continue
        candidates.append(
            {
                "chunk_id": int(row.chunk_id),
                "document_id": int(row.document_id),
                "document_title": row.title,
                "document_type": row.document_type,
                "source_ref": row.source_ref,
                "page_number": int(row.page_number) if row.page_number else None,
                "excerpt": _candidate_excerpt(row.content, question),
                "score": score,
                "preview_url": f"/api/knowledge/documents/{row.document_id}/preview",
            }
        )
    candidates.sort(key=lambda item: (-item["score"], item["document_id"], item["chunk_id"]))
    return candidates[:limit]


async def get_question_candidate_evidence(
    session: AsyncSession,
    question_id: int,
    *,
    limit: int = 5,
) -> list[dict] | None:
    result = await session.execute(
        text("SELECT question_text FROM knowledge_questions WHERE id = :question_id"),
        {"question_id": question_id},
    )
    question = result.one_or_none()
    if question is None:
        return None
    return await find_candidate_evidence(session, question.question_text, limit=limit)
